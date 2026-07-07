"""Tests for OUT-02 + Lock #5 DOCX renderer (Phase 5 Wave 2).

Covers:
  - register_index_styles: 4 styles registered (IndexHeading/Entry/Subentry/Subsubentry)
  - freeze_docx: RESEARCH §H-3 verbatim algorithm
    * FROZEN_DT == (1980,1,1,0,0,0); FROZEN_TS == "1970-01-01T00:00:00Z"
    * date_time of every zip entry is FROZEN_DT post-freeze
    * <dcterms:created> + <dcterms:modified> normalized to FROZEN_TS
  - **Byte-identity Lock #5 test (slow):** 2-second-gap saves byte-identical post-freeze
  - **No-XE Path A invariant (RESEARCH §H-9):** zero <w:fldSimple>/<w:fldChar>/<w:instrText>/`XE ` markers
  - Italic *(also: ...)* run emits <w:rPr><w:i/></w:rPr>
  - Metadata embedded in core.xml (dc:description) survives freeze
  - Sub-entries use IndexSubentry style; synthetic siblings use IndexSubentry
  - IndexSubsubentry registered but not applied on the reference corpus v1.0
  - Variant filter parity with markdown.py
  - FreezeError on corrupt input zip
  - Integration smoke (skipif live IR absent): output ≥30 KB

requirements_addressed: OUT-02, OUT-04 (Lock #5 byte-identity).
"""
from __future__ import annotations

import json
import re
import time
import zipfile
from pathlib import Path

import pytest

from book_indexer.render import (
    IndexEntry,
    IndexTree,
    Locator,
    SubEntry,
    SyntheticEntry,
)
from book_indexer.render.errors import FreezeError
from book_indexer.tables.ir import (
    CaseEntry,
    RuleEntry,
    StatuteEntry,
    SubsectionEntry,
    TableOfCases,
    TableOfRules,
    TableOfStatutes,
    TableProvenance,
)


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------


def _empty_tables() -> dict:
    prov = TableProvenance(
        eyecite_version="2.7.6",
        reporters_db_version="3.2.64",
        courts_db_version="0.10.27",
        pdf_sha256="a" * 64,
        corpus_sha="b" * 64,
        jurisdictions_enabled=[],
        chapter_rule_systems={},
        cite_counts={},
        regex_fallback_counts={},
        unresolved_short_cites=[],
        unverified_extractions=[],
    )
    return {
        "cases": TableOfCases(schema_version="1", entries=[], provenance=prov),
        "statutes": TableOfStatutes(schema_version="1", entries=[], provenance=prov),
        "rules": TableOfRules(schema_version="1", entries=[], provenance=prov),
    }


# --------------------------------------------------------------------------
# Module exports & constants
# --------------------------------------------------------------------------


def test_module_exports_required_names():
    from book_indexer.render import docx as d

    for name in (
        "render_docx",
        "register_index_styles",
        "freeze_docx",
        "FROZEN_DT",
        "FROZEN_TS",
    ):
        assert hasattr(d, name), f"docx.py missing export: {name}"


def test_frozen_constants():
    from book_indexer.render.docx import FROZEN_DT, FROZEN_TS

    assert FROZEN_DT == (1980, 1, 1, 0, 0, 0)
    assert FROZEN_TS == "1970-01-01T00:00:00Z"


# --------------------------------------------------------------------------
# Style registration (RESEARCH §H-2)
# --------------------------------------------------------------------------


def test_register_index_styles_adds_four_styles(tmp_path):
    from docx import Document

    from book_indexer.render.docx import register_index_styles

    d = Document()
    register_index_styles(d)
    style_names = {s.name for s in d.styles}
    assert "IndexHeading" in style_names
    assert "IndexEntry" in style_names
    assert "IndexSubentry" in style_names
    assert "IndexSubsubentry" in style_names

    # Save and verify all 4 names appear in word/styles.xml.
    out = tmp_path / "styles.docx"
    d.save(out)
    with zipfile.ZipFile(out, "r") as zf:
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
    for name in ("IndexHeading", "IndexEntry", "IndexSubentry", "IndexSubsubentry"):
        assert name in styles_xml, f"{name} not present in word/styles.xml"


# --------------------------------------------------------------------------
# freeze_docx — algorithm contract
# --------------------------------------------------------------------------


def _build_simple_docx(out_path: Path) -> None:
    """Build a 5-paragraph DOCX with the Phase 5 styles for freeze testing."""
    from docx import Document

    from book_indexer.render.docx import register_index_styles

    d = Document()
    register_index_styles(d)
    d.add_paragraph("Subject Index", style="IndexHeading")
    d.add_paragraph("admissibility, § 2.04 (p. 19)", style="IndexEntry")
    p = d.add_paragraph(style="IndexEntry")
    p.add_run("foo")
    italic = p.add_run(" *(also: bar)*")
    italic.italic = True
    p.add_run(", § 1.07 (p. 11)")
    d.add_paragraph("    sub-entry, § 3.02 (p. 50)", style="IndexSubentry")
    d.save(out_path)


def test_freeze_docx_zip_entry_dates_normalized(tmp_path):
    from book_indexer.render.docx import FROZEN_DT, freeze_docx

    raw = tmp_path / "raw.docx"
    _build_simple_docx(raw)
    frozen = tmp_path / "frozen.docx"
    freeze_docx(raw, frozen)

    with zipfile.ZipFile(frozen, "r") as zf:
        for info in zf.infolist():
            assert info.date_time == FROZEN_DT, (
                f"{info.filename} not normalized: {info.date_time}"
            )


def test_freeze_docx_dcterms_normalized(tmp_path):
    from book_indexer.render.docx import FROZEN_TS, freeze_docx

    raw = tmp_path / "raw.docx"
    _build_simple_docx(raw)
    frozen = tmp_path / "frozen.docx"
    freeze_docx(raw, frozen)

    with zipfile.ZipFile(frozen, "r") as zf:
        core_xml = zf.read("docProps/core.xml").decode("utf-8")
    assert FROZEN_TS in core_xml
    # Both <dcterms:created> and <dcterms:modified> should be normalized.
    created = re.search(r"<dcterms:created[^>]*>([^<]+)</dcterms:created>", core_xml)
    modified = re.search(r"<dcterms:modified[^>]*>([^<]+)</dcterms:modified>", core_xml)
    assert created and created.group(1) == FROZEN_TS
    assert modified and modified.group(1) == FROZEN_TS


@pytest.mark.slow
def test_freeze_docx_byte_identical_across_2s_gap(tmp_path):
    """RESEARCH §H-3 Lock #5 anchor — the load-bearing test for OUT-04.

    Build the same DOCX twice with a 2-second gap; freeze both; assert
    byte-identical. Mirrors the §H-3 empirical probe (sha1 anchor
    18a938b75e... in the research doc).
    """
    from book_indexer.render.docx import freeze_docx

    raw1 = tmp_path / "raw1.docx"
    raw2 = tmp_path / "raw2.docx"
    _build_simple_docx(raw1)
    time.sleep(2.1)
    _build_simple_docx(raw2)

    frozen1 = tmp_path / "frozen1.docx"
    frozen2 = tmp_path / "frozen2.docx"
    freeze_docx(raw1, frozen1)
    freeze_docx(raw2, frozen2)

    b1 = frozen1.read_bytes()
    b2 = frozen2.read_bytes()
    assert b1 == b2, (
        f"Lock #5 violation: frozen DOCXs differ ({len(b1)} vs {len(b2)} bytes)"
    )


def test_freeze_docx_raises_freeze_error_on_corrupt_input(tmp_path):
    from book_indexer.render.docx import freeze_docx

    bogus = tmp_path / "bogus.docx"
    bogus.write_bytes(b"this is not a zip file at all")
    out = tmp_path / "out.docx"
    with pytest.raises(FreezeError):
        freeze_docx(bogus, out)


# --------------------------------------------------------------------------
# render_docx end-to-end: no-XE invariant + structural assertions
# --------------------------------------------------------------------------


def test_render_docx_writes_file(
    make_entry, make_locator, make_provenance, frozen_metadata, tmp_path
):
    from book_indexer.render.docx import render_docx

    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[make_entry(canonical="admissibility")],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_render_docx_no_xe_field_markers(
    make_entry, make_locator, make_provenance, frozen_metadata, tmp_path
):
    """RESEARCH §H-9 Path A invariant — Path A by-construction should never
    emit <w:fldSimple>, <w:fldChar>, <w:instrText>, or 'XE '."""
    from book_indexer.render.docx import render_docx

    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[
            make_entry(canonical="admissibility"),
            make_entry(canonical="hearsay", variants=["hearsay-rule"]),
        ],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out)

    XE_MARKERS = ("fldSimple", "fldChar", "<w:instrText", "XE ")
    with zipfile.ZipFile(out, "r") as zf:
        for info in zf.infolist():
            if not info.filename.endswith(".xml"):
                continue
            content = zf.read(info.filename).decode("utf-8", errors="replace")
            for marker in XE_MARKERS:
                assert marker not in content, (
                    f"Path A violation: {marker!r} found in {info.filename}"
                )


def test_render_docx_italic_run_xml(
    make_entry, make_locator, make_provenance, frozen_metadata, tmp_path
):
    """An entry with variants emits <w:rPr><w:i/></w:rPr> on the variant run."""
    from book_indexer.render.docx import render_docx

    e = make_entry(canonical="foo", variants=["bar"])
    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[e],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out)

    with zipfile.ZipFile(out, "r") as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    # The italic flag is emitted as <w:rPr><w:i/></w:rPr> (or with attributes).
    # Check for the <w:i/> marker inside an <w:rPr> block.
    assert "<w:i" in doc_xml, "italic <w:i/> marker not found in document.xml"
    assert "(also: bar)" in doc_xml, "variant text 'bar' not in document.xml"


def test_render_docx_metadata_in_core_xml(
    make_entry, make_provenance, frozen_metadata, tmp_path
):
    """RESEARCH §H-11 Approach 1: metadata serialized into core.xml fields
    (e.g., <dc:description>). Verify pdf_sha256 fingerprint is present
    after freeze."""
    from book_indexer.render.docx import render_docx

    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[make_entry(canonical="admissibility")],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out)

    with zipfile.ZipFile(out, "r") as zf:
        core_xml = zf.read("docProps/core.xml").decode("utf-8")
    # The pdf_sha256 fingerprint is uniquely identifying.
    assert frozen_metadata.pdf_sha256 in core_xml, (
        "metadata pdf_sha256 not embedded in core.xml"
    )


def test_render_docx_subentry_uses_indexsubentry_style(
    make_entry, make_locator, make_provenance, frozen_metadata, tmp_path
):
    """Sub-entries must use the IndexSubentry style, not IndexEntry."""
    from docx import Document

    from book_indexer.render.docx import render_docx

    sub = SubEntry(
        text="under cross",
        sort_key="under cross",
        locators=[make_locator(section_ref="§3.02", folio="50")],
    )
    e = make_entry(canonical="hearsay", sub_entries=[sub])
    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[e],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out)

    d = Document(str(out))
    sub_para = next(
        (p for p in d.paragraphs if "under cross" in p.text),
        None,
    )
    assert sub_para is not None, "sub-entry paragraph not found"
    assert sub_para.style.name == "IndexSubentry"


def test_render_docx_synthetic_uses_correct_styles(
    make_locator, make_provenance, frozen_metadata, tmp_path
):
    """Synthetic stem uses IndexEntry; sibling lines use IndexSubentry."""
    from docx import Document

    from book_indexer.render.docx import render_docx

    syn = SyntheticEntry(
        stem="hearsay",
        sibling_canonicals=("admissible hearsay",),
        locators=(make_locator(section_ref="§2.04", folio="19"),),
    )
    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [syn], _empty_tables(), frozen_metadata, out)

    d = Document(str(out))
    # Phase 7 CUR-03 sub-rule b: synthesized parent emits bare stem (no marker).
    # Match on the bare stem text on a paragraph styled IndexEntry.
    head = next(
        (
            p
            for p in d.paragraphs
            if p.text == "hearsay" and p.style.name == "IndexEntry"
        ),
        None,
    )
    sib = next(
        (p for p in d.paragraphs if "admissible hearsay" in p.text), None
    )
    assert head is not None
    assert "(synthesized)" not in head.text
    assert sib is not None
    assert sib.style.name == "IndexSubentry"


def test_render_docx_indexsubsubentry_registered_but_unused(
    make_entry, make_provenance, frozen_metadata, tmp_path
):
    """IndexSubsubentry is registered but no paragraph uses it on the reference corpus v1.0."""
    from docx import Document

    from book_indexer.render.docx import render_docx

    tree = IndexTree(
        schema_version="1",
        provenance=make_provenance(),
        entries=[make_entry(canonical="admissibility")],
    )
    out = tmp_path / "out.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out)

    d = Document(str(out))
    style_names = {s.name for s in d.styles}
    assert "IndexSubsubentry" in style_names
    used = {p.style.name for p in d.paragraphs}
    assert "IndexSubsubentry" not in used


# --------------------------------------------------------------------------
# Variant filter parity with markdown.py
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "canonical,variants,expected",
    [
        ("foo", ["Foo", "FOO"], []),  # case-only drop
        ("foo", ["(*foo", "[bracket]", "ok"], ["ok"]),  # cruft drop
        ("foo", ["a", "bb", "cccc", "ddd", "eeee"], ["cccc", "eeee", "ddd"]),  # top-3
        ("foo", ["bar"], ["bar"]),  # single variant survives
    ],
)
def test_filter_variants_parity_with_markdown(canonical, variants, expected):
    from book_indexer.render.docx import _filter_variants as docx_filter
    from book_indexer.render.markdown import _filter_variants as md_filter

    md_out = md_filter(canonical, variants)
    docx_out = docx_filter(canonical, variants)
    assert md_out == docx_out
    assert md_out == expected


# --------------------------------------------------------------------------
# Determinism — render_docx twice should produce byte-identical output
# (after freeze; this is the load-bearing OUT-04 anchor for DOCX)
# --------------------------------------------------------------------------


@pytest.mark.slow
def test_render_docx_deterministic_2s_gap(
    make_entry, make_locator, make_provenance, frozen_metadata, tmp_path
):
    """End-to-end Lock #5: render_docx twice with a 2-second gap; assert
    byte-identical output."""
    from book_indexer.render.docx import render_docx

    e = make_entry(
        canonical="hearsay",
        variants=["hearsay rule"],
        locators=[make_locator(section_ref="§2.04", folio="19")],
    )
    tree = IndexTree(
        schema_version="1", provenance=make_provenance(), entries=[e]
    )

    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"
    render_docx(tree, [], _empty_tables(), frozen_metadata, out1)
    time.sleep(2.1)
    render_docx(tree, [], _empty_tables(), frozen_metadata, out2)

    assert out1.read_bytes() == out2.read_bytes(), (
        "Lock #5 violation: render_docx not deterministic across 2s gap"
    )


# --------------------------------------------------------------------------
# Integration smoke (skipif live IR absent)
# --------------------------------------------------------------------------


_LIVE_IR = Path("artifacts/index_tree.json")
_LIVE_TABLES_DIR = Path("artifacts/tables")


@pytest.mark.slow
@pytest.mark.skipif(
    not _LIVE_IR.exists() or not (_LIVE_TABLES_DIR / "cases.json").exists(),
    reason="live IR / tables absent",
)
def test_integration_smoke_live_ir(frozen_metadata, tmp_path):
    """Full DOCX render against live IR + tables ≥ 30 KB (RESEARCH §H-12)."""
    import spacy

    from book_indexer.render.docx import render_docx
    from book_indexer.render.filter import is_cruft
    from book_indexer.render.synthesize import (
        load_stopwords,
        synthesize_bare_lemma_entries,
    )

    tree = IndexTree.model_validate(json.loads(_LIVE_IR.read_text()))
    cases = TableOfCases.model_validate(
        json.loads((_LIVE_TABLES_DIR / "cases.json").read_text())
    )
    statutes = TableOfStatutes.model_validate(
        json.loads((_LIVE_TABLES_DIR / "statutes.json").read_text())
    )
    rules = TableOfRules.model_validate(
        json.loads((_LIVE_TABLES_DIR / "rules.json").read_text())
    )
    tables = {"cases": cases, "statutes": statutes, "rules": rules}

    surviving = [e for e in tree.entries if not is_cruft(e.canonical)]
    nlp = spacy.load("en_core_web_lg")
    stopwords = load_stopwords()
    synthetics = synthesize_bare_lemma_entries(surviving, nlp, stopwords)

    out = tmp_path / "live.docx"
    render_docx(tree, synthetics, tables, frozen_metadata, out)
    assert out.stat().st_size >= 30_000, (
        f"live DOCX render too small: {out.stat().st_size} bytes"
    )
