"""OUT-05 + Lock #5 — DOCX sections-only renderer (Phase 7 Wave 3).

Near-copy of ``docx.py`` (CONTEXT D-02 — "renderer near-copy boundary")
that emits the *sections-only* DOCX variant. Two structural deltas only:

  1. Locator formatter swap: ``collapse_locators`` →
     ``collapse_locators_sections_only`` (D-03 sections-only collapse;
     emits ``§ N.NN`` / ``§§ N–M`` only).
  2. ``Metadata.pages_only_variant`` is set to ``True`` for this variant.

REUSED UNCHANGED from docx.py (Lock #5 byte-determinism — DO NOT
duplicate or fork these): ``XE_MARKERS``, ``FROZEN_DT``, ``FROZEN_TS``,
``_RE_CREATED``, ``_RE_MODIFIED``, ``freeze_docx``,
``register_index_styles``. The freeze post-process is the load-bearing
byte-identity contract.

Wave 3 Task 2 then extends this file with curator-overrides wiring
(removal short-circuit + recapitalize pre-emit + CUR-03 plural-variant
filter + (synthesized)-marker drop). Pitfall §P-3 (Pitfall 3 of
07-RESEARCH): recapitalize is applied at the STRING level BEFORE
``add_run(text)`` so DOCX run-segmentation never breaks the exact-string
substitution.

requirements_addressed: OUT-05, OUT-04 (via ``freeze_docx`` reuse).
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

from docx import Document

from book_indexer.curator import (
    CuratorOverrides,
    apply_recap_pairs,
    assert_letters_only,
    is_droppable_plural_variant,
)
from book_indexer.curator.fixture import EditorialOverrides

from .cross_refs import derive_cross_refs
from .docx import (
    FROZEN_DT,
    FROZEN_TS,
    XE_MARKERS,
    _RE_CREATED,
    _RE_MODIFIED,
    _recap,
    _render_cross_ref,
    freeze_docx,
    register_index_styles,
)
from .editorial_overrides import (
    apply_editorial_overrides,
    apply_r9_whitespace,
    write_mismatch_report,
)
from .filter import is_cruft, is_removed
from .ir import IndexEntry, IndexTree, SubEntry, SyntheticEntry
from .markdown_sections_only import _dedup_statute_entries
from .metadata import Metadata
from .parent_dedup import dedupe_parent_aliased_standalones
from .plural_consolidation import (
    ConsolidatedEntry,
    consolidate_plural_pairs,
)
from .section_range_collapse import collapse_locators_sections_only
from book_indexer.tables.ir import (
    Locator,
    TableOfCases,
    TableOfRules,
    TableOfStatutes,
)

__all__ = [
    "render_docx_sections_only",
]


# Locator-list separator inside an entry line.
_LOCATOR_SEP = ", "
_VARIANT_SEP = "; "
_MAX_VARIANTS = 3


# --------------------------------------------------------------------------
# Variant filter — mirrors docx.py contract; Task 2 extends with
# is_droppable_plural_variant + keep_plural_set kwarg.
# --------------------------------------------------------------------------


def _filter_variants(
    canonical: str,
    variants: list[str],
    keep_plural_set: frozenset[str] = frozenset(),
    dropped_plural_acc: list[tuple[str, str]] | None = None,
) -> list[str]:
    """Drop case-only + cruft-shaped + (CUR-03) plural variants; sort + top-3.

    Identical contract to ``docx.py:_filter_variants`` (Phase 7 Wave 3
    extension). CUR-03 applies to BOTH output variants per CONTEXT D-13.
    """
    canon_lower = canonical.lower()
    survivors: list[str] = []
    seen: set[str] = set()
    for v in variants:
        if v.lower() == canon_lower:
            continue
        if is_cruft(v):
            continue
        if is_droppable_plural_variant(canonical, v, keep_plural_set):
            if dropped_plural_acc is not None:
                dropped_plural_acc.append((canonical, v))
            continue
        key = v.lower()
        if key in seen:
            continue
        seen.add(key)
        survivors.append(v)
    survivors.sort(key=lambda v: (-len(v), v))
    return survivors[:_MAX_VARIANTS]


# --------------------------------------------------------------------------
# Locator helpers — sections-only formatter
# --------------------------------------------------------------------------


def _render_locators_string(locators: list[Locator]) -> str:
    """OUT-05: emit only § N.NN (no (p. N)); apply D-03 sections-only collapse."""
    formatted = collapse_locators_sections_only(locators)
    return _LOCATOR_SEP.join(fl.rendered for fl in formatted)


# --------------------------------------------------------------------------
# Metadata embedding — mirrors docx.py
# --------------------------------------------------------------------------


def _serialize_metadata(metadata: Metadata) -> str:
    payload = metadata.model_dump()
    payload.pop("built_at", None)
    lines = ["book-indexer metadata"]
    for key in sorted(payload):
        lines.append(f"{key}={payload[key]}")
    return "\n".join(lines)


def _render_metadata_into_doc(d: Document, metadata: Metadata) -> None:
    """Mirrors docx.py — embeds pdf_sha256 + serialized block on the doc."""
    d.core_properties.subject = metadata.pdf_sha256
    d._tb_metadata_block = _serialize_metadata(metadata)  # type: ignore[attr-defined]


def _emit_metadata_block_paragraph(d: Document) -> None:
    block = getattr(d, "_tb_metadata_block", None)
    if not block:
        return
    d.add_paragraph("Metadata", style="IndexHeading")
    for line in block.split("\n"):
        d.add_paragraph(line, style="IndexEntry")


# --------------------------------------------------------------------------
# Entry / sub-entry / synthetic / table renderers (sections-only locators)
# --------------------------------------------------------------------------


def _render_entry(
    d: Document,
    entry: IndexEntry,
    pairs: list[tuple[str, str]] | None = None,
    keep_plural_set: frozenset[str] = frozenset(),
    dropped_plural_acc: list[tuple[str, str]] | None = None,
    removal_set: frozenset[str] = frozenset(),
) -> None:
    """Wave 4 CUR-01 sub-entry inheritance: drop subs whose ``text`` is in
    ``removal_set`` (mirrors ``docx.py:_render_entry`` Wave 4 fix)."""
    p = d.add_paragraph(style="IndexEntry")
    p.add_run(_recap(entry.canonical, pairs))

    variants = _filter_variants(
        entry.canonical, entry.variants, keep_plural_set, dropped_plural_acc
    )
    if variants:
        also_text = " *(also: " + _VARIANT_SEP.join(variants) + ")*"
        italic = p.add_run(_recap(also_text, pairs))
        italic.italic = True

    locator_str = _render_locators_string(entry.locators)
    if locator_str:
        p.add_run(_recap(f", {locator_str}", pairs))

    for sub in sorted(entry.sub_entries, key=lambda s: s.sort_key):
        if removal_set and is_removed(sub.text, removal_set):
            continue
        sub_p = d.add_paragraph(style="IndexSubentry")
        sub_p.add_run(_recap(sub.text, pairs))
        sub_loc = _render_locators_string(sub.locators)
        if sub_loc:
            sub_p.add_run(_recap(f", {sub_loc}", pairs))


def _render_synthetic(
    d: Document,
    synth: SyntheticEntry,
    entries_by_canonical: dict[str, IndexEntry],
    pairs: list[tuple[str, str]] | None = None,
    removal_set: frozenset[str] = frozenset(),
    transferred_variants: dict[str, tuple[str, ...]] | None = None,
) -> None:
    """Synthesized parent + sibling sub-entries (sections-only locators).

    CUR-03 sub-rule b: literal ``(synthesized)`` marker NOT emitted; the
    IR field ``IndexEntry.synthesized`` (and ``SyntheticEntry`` itself)
    is unchanged.

    Wave 4 CUR-01 sub-rule c: sibling canonicals matching curator-removed
    terms are dropped from the synthetic block.

    v1.2.2: ``transferred_variants`` carries the dedup pass's side-channel
    map ``{child_canonical: variants}`` — emit italic ``(also: …)`` runs
    on surviving children whose standalone was dropped (mirrors
    ``docx.py:_render_synthetic``).
    """
    transferred_variants = transferred_variants or {}
    head = d.add_paragraph(style="IndexEntry")
    head.add_run(_recap(synth.stem, pairs))

    for sib_canonical in sorted(synth.sibling_canonicals):
        if removal_set and is_removed(sib_canonical, removal_set):
            continue
        sib_entry = entries_by_canonical.get(sib_canonical)
        if sib_entry is not None:
            loc_str = _render_locators_string(sib_entry.locators)
        else:
            loc_str = _render_locators_string(list(synth.locators))
        sub_p = d.add_paragraph(style="IndexSubentry")
        sub_p.add_run(_recap(sib_canonical, pairs))
        # v1.2.2: optional transferred variants — italic (also: ...) run
        # mirroring the standalone entry's _render_entry rendering.
        variants = transferred_variants.get(sib_canonical, ())
        if variants:
            filtered = _filter_variants(sib_canonical, list(variants))
            if filtered:
                also_text = " *(also: " + _VARIANT_SEP.join(filtered) + ")*"
                italic = sub_p.add_run(_recap(also_text, pairs))
                italic.italic = True
        if loc_str:
            sub_p.add_run(_recap(f", {loc_str}", pairs))


def _entry_letter(sort_key: str) -> str:
    return sort_key[0].upper() if sort_key else ""


def _render_consolidated(
    d: Document,
    consolidated: ConsolidatedEntry,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    """v1.2.1 sections-only: emit a ConsolidatedEntry as an IndexEntry
    paragraph using the sections-only locator formatter.

    Mirrors ``docx.py:_render_consolidated`` shape; only delta is the
    locator formatter (D-03 sections-only).
    """
    p = d.add_paragraph(style="IndexEntry")
    if consolidated.source_kind == "xref":
        text = (
            f"{consolidated.display_canonical}. "
            f"See {consolidated.see_target}."
        )
        p.add_run(_recap(text, pairs))
        return
    p.add_run(_recap(consolidated.display_canonical, pairs))
    locator_str = _render_locators_string(list(consolidated.locators))  # type: ignore[arg-type]
    if locator_str:
        p.add_run(_recap(f", {locator_str}", pairs))


def _render_subject_index(
    d: Document,
    tree: IndexTree,
    synthetics: list[SyntheticEntry],
    pairs: list[tuple[str, str]] | None = None,
    removal_set: frozenset[str] = frozenset(),
    keep_plural_set: frozenset[str] = frozenset(),
    dropped_plural_acc: list[tuple[str, str]] | None = None,
    editorial_overrides: EditorialOverrides | None = None,
    allow_stale_overrides: bool = False,
) -> None:
    d.add_paragraph(_recap("Subject Index", pairs), style="IndexHeading")

    surviving = [
        e
        for e in tree.entries
        if not is_cruft(e.canonical) and not is_removed(e.canonical, removal_set)
    ]

    # Phase 9 — 7th projection pass (mirrors markdown.py / docx.py).
    _xref_removal_set: frozenset[str] = frozenset()
    _synth_suppressed_stems: frozenset[str] = frozenset()
    if editorial_overrides is not None:
        _apply_result = apply_editorial_overrides(
            surviving,
            editorial_overrides,
            allow_stale=allow_stale_overrides,
        )
        surviving = _apply_result.entries
        _xref_removal_set = _apply_result.xref_removal_set
        _synth_suppressed_stems = _apply_result.synth_suppressed_stems
        write_mismatch_report(
            _apply_result.mismatches,
            Path("artifacts/audit/editorial_override_mismatches.md"),
        )

    entries_by_canonical = {e.canonical: e for e in surviving}

    # UAT 08-1 head-noun cross-refs (mirrors markdown_sections_only.py).
    # Phase 9: R6 filters synth_suppressed_stems; R5 filters cross-refs.
    surviving_synthetics = [
        s for s in synthetics
        if not is_removed(s.stem, removal_set)
        and s.stem not in _synth_suppressed_stems
    ]
    cross_refs = derive_cross_refs(surviving, surviving_synthetics)
    if _xref_removal_set:
        cross_refs = [x for x in cross_refs if x.head not in _xref_removal_set]

    sorted_entries = sorted(surviving, key=lambda e: e.sort_key)
    sorted_synthetics = sorted(surviving_synthetics, key=lambda s: s.stem)

    merged: list[tuple[str, str, object]] = []
    for e in sorted_entries:
        merged.append((e.sort_key, "entry", e))
    for s in sorted_synthetics:
        merged.append((s.stem.lower(), "synth", s))
    for x in cross_refs:
        merged.append((x.sort_key, "xref", x))
    merged.sort(key=lambda t: t[0])

    # v1.2.1: collapse adjacent singular/plural pairs (mirrors markdown.py).
    merged = consolidate_plural_pairs(merged, keep_plural_set)

    # v1.2.2: parent-aliased-standalone dedup (mirrors markdown.py).
    merged, transferred_variants, _dedup_decisions = (
        dedupe_parent_aliased_standalones(merged, entries_by_canonical)
    )

    current_letter = ""
    for sort_key, kind, payload in merged:
        letter = _entry_letter(sort_key)
        if letter != current_letter:
            current_letter = letter
            d.add_paragraph(letter, style="IndexHeading")
        if kind == "entry":
            _render_entry(
                d,
                payload,  # type: ignore[arg-type]
                pairs,
                keep_plural_set,
                dropped_plural_acc,
                removal_set,
            )
        elif kind == "synth":
            _render_synthetic(
                d,
                payload,  # type: ignore[arg-type]
                entries_by_canonical,
                pairs,
                removal_set,
                transferred_variants,
            )
        elif kind == "consolidated":
            _render_consolidated(d, payload, pairs)  # type: ignore[arg-type]
        else:  # xref — UAT 08-1
            _render_cross_ref(d, payload, pairs)  # type: ignore[arg-type]


def _render_table_of_cases(
    d: Document,
    tables: TableOfCases | None,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    d.add_paragraph(_recap("Table of Cases", pairs), style="IndexHeading")
    if tables is None:
        return
    for case in sorted(tables.entries, key=lambda c: c.sort_key):
        p = d.add_paragraph(style="IndexEntry")
        name_run = p.add_run(_recap(case.display_name, pairs))
        name_run.italic = True
        p.add_run(_recap(f", {case.canonical_citation}", pairs))
        loc_str = _render_locators_string(case.locators)
        if loc_str:
            p.add_run(_recap(f", {loc_str}", pairs))


def _render_table_of_statutes(
    d: Document,
    tables: TableOfStatutes | None,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    d.add_paragraph(_recap("Table of Statutes", pairs), style="IndexHeading")
    if tables is None:
        return
    deduped = _dedup_statute_entries(tables.entries)
    for s in sorted(deduped, key=lambda x: x.sort_key):
        p = d.add_paragraph(style="IndexEntry")
        p.add_run(_recap(s.display_name, pairs))
        loc_str = _render_locators_string(s.locators)
        if loc_str:
            p.add_run(_recap(f", {loc_str}", pairs))


def _render_table_of_rules(
    d: Document,
    tables: TableOfRules | None,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    d.add_paragraph(_recap("Table of Rules", pairs), style="IndexHeading")
    if tables is None:
        return
    for rule in sorted(tables.entries, key=lambda r: r.sort_key):
        p = d.add_paragraph(style="IndexEntry")
        p.add_run(_recap(rule.parent_rule, pairs))
        parent_loc = _render_locators_string(rule.parent_locators)
        if parent_loc:
            p.add_run(_recap(f", {parent_loc}", pairs))
        for sub in sorted(rule.subsections, key=lambda s: s.subsection_path):
            sub_p = d.add_paragraph(style="IndexSubentry")
            sub_p.add_run(_recap(f"{rule.parent_rule}{sub.subsection_path}", pairs))
            sub_loc = _render_locators_string(sub.locators)
            if sub_loc:
                sub_p.add_run(_recap(f", {sub_loc}", pairs))


# --------------------------------------------------------------------------
# Public API
# --------------------------------------------------------------------------


def render_docx_sections_only(
    tree: IndexTree,
    synthetics: list[SyntheticEntry],
    tables: dict[str, object],
    metadata: Metadata,
    out_path: Path,
    overrides: CuratorOverrides | None = None,
    editorial_overrides: EditorialOverrides | None = None,
    curator_log: dict[str, list] | None = None,
) -> None:
    """Render OUT-05 — write index_sections_only.docx to ``out_path``.

    Mirrors ``render_docx`` extension shape — same curator-overrides
    contract (CUR-01 + CUR-02 + CUR-03) applied to BOTH variants. Only
    delta vs. ``render_docx``: locator formatter (sections-only) +
    ``pages_only_variant=True``.
    """
    out_path = Path(out_path)
    cases_table = tables.get("cases")  # type: ignore[union-attr]
    statutes_table = tables.get("statutes")  # type: ignore[union-attr]
    rules_table = tables.get("rules")  # type: ignore[union-attr]

    metadata_for_block = metadata.model_copy(update={"pages_only_variant": True})

    pairs: list[tuple[str, str]] | None = None
    removal_set: frozenset[str] = frozenset()
    keep_plural_set: frozenset[str] = frozenset()
    dropped_plural_acc: list[tuple[str, str]] = []

    if overrides is not None:
        removal_set = overrides.removal_set
        keep_plural_set = overrides.keep_plural_set
        if overrides.recapitalize_pairs:
            pairs = [tuple(p) for p in overrides.recapitalize_pairs]
            assert_letters_only(pairs)

    # Phase 9 — D-07 ALLOW_STALE_OVERRIDES env-flag (mirrors markdown.py).
    allow_stale_overrides = (
        os.environ.get("ALLOW_STALE_OVERRIDES") == "1"
    )

    d = Document()
    register_index_styles(d)
    _render_metadata_into_doc(d, metadata_for_block)

    _render_subject_index(
        d, tree, synthetics, pairs, removal_set, keep_plural_set, dropped_plural_acc,
        editorial_overrides=editorial_overrides,
        allow_stale_overrides=allow_stale_overrides,
    )
    _render_table_of_cases(d, cases_table, pairs)  # type: ignore[arg-type]
    _render_table_of_statutes(d, statutes_table, pairs)  # type: ignore[arg-type]
    _render_table_of_rules(d, rules_table, pairs)  # type: ignore[arg-type]
    _emit_metadata_block_paragraph(d)

    # Phase 9 — R9 whitespace at run-text level (mirrors docx.py).
    if editorial_overrides is not None and editorial_overrides.R9_whitespace:
        r9_rules = editorial_overrides.R9_whitespace
        for paragraph in d.paragraphs:
            for run in paragraph.runs:
                run.text = apply_r9_whitespace(run.text, r9_rules)

    if curator_log is not None and overrides is not None:
        curator_log.setdefault("dropped_plural_variants", []).extend(
            dropped_plural_acc
        )

    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, dir=out_path.parent
    ) as tmp_fp:
        tmp_path = Path(tmp_fp.name)
    try:
        d.save(str(tmp_path))
        freeze_docx(tmp_path, out_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()
