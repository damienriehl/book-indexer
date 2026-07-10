"""OUT-02 + Lock #5 — DOCX index renderer (Phase 5 Wave 2).

Emits the user-facing `index.docx` deliverable via python-docx 1.2.0
**baked-paragraphs (Path A)** + the RESEARCH §H-3 verbatim `freeze_docx`
post-process for Lock #5 byte-identical determinism across arbitrary
time gaps.

Architecture locks enforced here:

  Lock #5 (byte-identity) — every output byte-identical across runs.
    Implementation: sorted iterations + frozen Unix-epoch zip date_time +
    regex-normalized <dcterms:created>/<dcterms:modified> in core.xml.
    Verified empirically in §H-3 (sha1 18a938b75e... matched across a
    2-second gap).

  Path A (no XE fields) — RESEARCH §H-9 invariant. python-docx's
    `add_paragraph(style=...)` + `paragraph.add_run(...)` API surface
    NEVER emits <w:fldSimple>, <w:fldChar>, <w:instrText>, or 'XE '
    markers. We use ONLY this surface — no `add_field()` calls
    anywhere — so the no-XE property is by-construction. The unit
    test `test_render_docx_no_xe_field_markers` enforces it.

D-02 — italic case-derived concepts in subject index DEFERRED to v1.x.
`IndexEntry.derived_from_table` is IGNORED here.

Pitfall §P-2 — preserve zip entry order (use `zin.infolist()` order).
Pitfall §P-5 — namespace drift in core.xml: regex matches the leading
`<dcterms:created` / `<dcterms:modified` element start tag with optional
attributes; we substitute the full open+close including the W3CDTF
xsi:type attribute (matches python-docx 1.2.0 default emit).

requirements_addressed: OUT-02, OUT-04 (Lock #5 byte-identity).
"""
from __future__ import annotations

import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import cast

from docx import Document as _new_docx_document
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle

from book_indexer.curator import (
    CuratorOverrides,
    apply_recap_pairs,
    assert_letters_only,
    is_droppable_plural_variant,
)
from book_indexer.curator.fixture import EditorialOverrides
from book_indexer.tables.ir import (
    Locator,
    TableOfCases,
    TableOfRules,
    TableOfStatutes,
)

from .cross_refs import CrossRefEntry, derive_cross_refs
from .editorial_overrides import (
    apply_editorial_overrides,
    apply_r9_whitespace,
    write_mismatch_report,
)
from .errors import FreezeError
from .filter import is_cruft, is_removed
from .ir import IndexEntry, IndexTree, SyntheticEntry
from .markdown import _dedup_statute_entries
from .metadata import Metadata
from .parent_dedup import dedupe_parent_aliased_standalones
from .plural_consolidation import (
    ConsolidatedEntry,
    consolidate_plural_pairs,
)
from .range_collapse import collapse_locators

__all__ = [
    "FROZEN_DT",
    "FROZEN_TS",
    "freeze_docx",
    "register_index_styles",
    "render_docx",
]


# --------------------------------------------------------------------------
# Constants — RESEARCH §H-3 verbatim
# --------------------------------------------------------------------------


FROZEN_DT: tuple[int, int, int, int, int, int] = (1980, 1, 1, 0, 0, 0)
FROZEN_TS: str = "1970-01-01T00:00:00Z"

_RE_CREATED = re.compile(r"<dcterms:created[^>]*>[^<]+</dcterms:created>")
_RE_MODIFIED = re.compile(r"<dcterms:modified[^>]*>[^<]+</dcterms:modified>")

# RESEARCH §H-9 — Path A by-construction. Any of these markers in word/*.xml
# would mean a field-API call snuck in. The unit test enforces zero matches.
XE_MARKERS: tuple[str, ...] = ("fldSimple", "fldChar", "<w:instrText", "XE ")


# Locator-list separator inside an entry line.
_LOCATOR_SEP = ", "
# Variant separator inside the (also: ...) parenthetical.
_VARIANT_SEP = "; "
# Maximum number of variants emitted in any (also: ...) parenthetical.
_MAX_VARIANTS = 3


# --------------------------------------------------------------------------
# Style registration — RESEARCH §H-2 verbatim
# --------------------------------------------------------------------------


def register_index_styles(d: Document) -> None:
    """Register the 4 Phase 5 paragraph styles on a python-docx Document.

    RESEARCH §H-2 verbatim. Styles:
      - IndexHeading (Times New Roman 12pt bold) — letter dividers + table titles
      - IndexEntry (Times New Roman 10pt regular, 0in indent) — main entries
      - IndexSubentry (Times New Roman 10pt, 0.25in indent) — sub-entries
      - IndexSubsubentry (Times New Roman 10pt, 0.5in indent) — forward-compat;
        REGISTERED BUT NOT APPLIED on the reference corpus.
    """
    s = d.styles

    h = cast(ParagraphStyle, s.add_style("IndexHeading", WD_STYLE_TYPE.PARAGRAPH))
    h.font.name = "Times New Roman"
    h.font.size = Pt(12)
    h.font.bold = True

    e = cast(ParagraphStyle, s.add_style("IndexEntry", WD_STYLE_TYPE.PARAGRAPH))
    e.font.name = "Times New Roman"
    e.font.size = Pt(10)
    e.paragraph_format.left_indent = Inches(0)

    sub = cast(ParagraphStyle, s.add_style("IndexSubentry", WD_STYLE_TYPE.PARAGRAPH))
    sub.font.name = "Times New Roman"
    sub.font.size = Pt(10)
    sub.paragraph_format.left_indent = Inches(0.25)

    sub2 = cast(ParagraphStyle, s.add_style("IndexSubsubentry", WD_STYLE_TYPE.PARAGRAPH))
    sub2.font.name = "Times New Roman"
    sub2.font.size = Pt(10)
    sub2.paragraph_format.left_indent = Inches(0.5)


# --------------------------------------------------------------------------
# freeze_docx — RESEARCH §H-3 verbatim (do NOT modify the algorithm)
# --------------------------------------------------------------------------


def freeze_docx(in_path: Path, out_path: Path) -> None:
    """Rewrite a python-docx-emitted .docx so two builds are byte-identical.

    Strategy (RESEARCH §H-3 verbatim):
      1. Read every entry from the input zip in original order
         (Pitfall §P-2 — never `os.listdir` + sort).
      2. For docProps/core.xml, regex-normalize <dcterms:created> and
         <dcterms:modified> to FROZEN_TS.
      3. Rewrite each entry into a new zip with date_time=FROZEN_DT.
      4. Preserve compress_type=ZIP_DEFLATED and external_attr.

    Verified byte-identical across a 2-second gap on production-shape DOCX
    (50 IndexEntry paragraphs, ~37 KB output). See RESEARCH §H-3 — frozen
    sha1 anchor: 18a938b75e03342100e9fd7df1c551bc71ca168a99dad8cd8944a286cc8476ba.

    Raises:
        FreezeError: if the input is not a valid zip / DOCX.
    """
    try:
        with zipfile.ZipFile(in_path, "r") as zin, \
             zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                content = zin.read(info.filename)
                if info.filename == "docProps/core.xml":
                    text = content.decode("utf-8")
                    text = _RE_CREATED.sub(
                        f'<dcterms:created xsi:type="dcterms:W3CDTF">'
                        f"{FROZEN_TS}</dcterms:created>",
                        text,
                    )
                    text = _RE_MODIFIED.sub(
                        f'<dcterms:modified xsi:type="dcterms:W3CDTF">'
                        f"{FROZEN_TS}</dcterms:modified>",
                        text,
                    )
                    content = text.encode("utf-8")
                new_info = zipfile.ZipInfo(info.filename, date_time=FROZEN_DT)
                new_info.compress_type = zipfile.ZIP_DEFLATED
                new_info.external_attr = info.external_attr
                zout.writestr(new_info, content)
    except (zipfile.BadZipFile, OSError, KeyError) as e:
        raise FreezeError(f"freeze_docx failed on {in_path}: {e}") from e


# --------------------------------------------------------------------------
# Variant filter — MUST be byte-identical-in-behavior to markdown.py's
# _filter_variants. Test parametrize enforces parity.
# --------------------------------------------------------------------------


def _filter_variants(
    canonical: str,
    variants: list[str],
    keep_plural_set: frozenset[str] = frozenset(),
    dropped_plural_acc: list[tuple[str, str]] | None = None,
) -> list[str]:
    """Drop case-only + cruft-shaped + (CUR-03) plural variants; sort + top-3.

    Identical contract to ``markdown.py:_filter_variants`` (Phase 7 Wave 3
    extension; cross-test parametrize enforces parity).
    """
    canon_lower = canonical.lower()
    survivors: list[str] = []
    seen: set[str] = set()
    for v in variants:
        if v.lower() == canon_lower:
            continue
        if is_cruft(v):
            continue
        if is_droppable_plural_variant(canonical, v, keep_plural_set):
            if dropped_plural_acc is not None:
                dropped_plural_acc.append((canonical, v))
            continue
        key = v.lower()
        if key in seen:
            continue
        seen.add(key)
        survivors.append(v)
    survivors.sort(key=lambda v: (-len(v), v))
    return survivors[:_MAX_VARIANTS]


# --------------------------------------------------------------------------
# Locator helpers
# --------------------------------------------------------------------------


def _render_locators_string(locators: list[Locator]) -> str:
    formatted = collapse_locators(locators)
    return _LOCATOR_SEP.join(fl.rendered for fl in formatted)


# --------------------------------------------------------------------------
# CUR-02 recapitalize — string-level transform applied BEFORE add_run() so
# DOCX run-segmentation (Pitfall 3 of 07-RESEARCH) can never fragment the
# exact-string substitution. The renderers call `_recap(text, pairs)` before
# every `paragraph.add_run(text)` callsite when overrides are active.
# --------------------------------------------------------------------------


def _recap(text: str, pairs: list[tuple[str, str]] | None) -> str:
    """Apply recapitalize ``pairs`` to ``text`` if pairs are provided.

    No-op when ``pairs`` is None or empty (the renderer-without-overrides
    path). The strict-guard validation runs once at the top of
    ``render_docx`` (defense in depth — Pydantic's record-level
    validator already gates fixture load).
    """
    if not pairs:
        return text
    return apply_recap_pairs(text, pairs)


# --------------------------------------------------------------------------
# Metadata embedding — RESEARCH §H-11 Approach 1 (core_properties.comments)
# --------------------------------------------------------------------------


def _serialize_metadata(metadata: Metadata) -> str:
    """Serialize metadata to a key=value newline block (mirrors markdown
    block, minus HTML wrapper). Embedded in core.xml's <dc:description> via
    core_properties.comments — survives freeze, round-trippable on re-load."""
    payload = metadata.model_dump()
    payload.pop("built_at", None)
    lines = ["book-indexer metadata"]
    for key in sorted(payload):
        lines.append(f"{key}={payload[key]}")
    return "\n".join(lines)


def _render_metadata_into_doc(d: Document, metadata: Metadata) -> None:
    """Embed the AUD-04 metadata into core.xml AND a body paragraph.

    python-docx 1.2.0 enforces a 255-char limit on every core_properties
    string field (subject, comments=description, keywords, title, ...).
    The full AUD-04 block runs ~400 chars, so we split:

      1. ``pdf_sha256`` (64 chars) → ``core_properties.subject``
         (rendered as ``<dc:subject>``). This is the load-bearing
         fingerprint for AUD-04 + Lock #5 replay verification, and
         satisfies the unit test `metadata in core.xml`.
      2. Full block → a body paragraph with the IndexEntry style at the
         very end of the document. Word-based reviewers see it; the
         freeze post-process leaves it untouched.

    Per plan deviations_allowed (subject path is permitted) + RESEARCH
    §H-11. The freeze post-process does NOT touch either location —
    only <dcterms:created> + <dcterms:modified> are regex-normalized.
    """
    # 64 chars — well under the 255 cap.
    d.core_properties.subject = metadata.pdf_sha256
    # Cache the full block on the document for the body-paragraph emit;
    # render_docx adds it last so it doesn't disrupt subject-index sort.
    d._tb_metadata_block = _serialize_metadata(metadata)  # type: ignore[attr-defined]


def _emit_metadata_block_paragraph(d: Document) -> None:
    """Append the full AUD-04 metadata block as a body paragraph at the
    very end of the document (audit footer). Each line of the serialized
    block becomes its own paragraph with the IndexEntry style for
    determinism + visibility."""
    block = getattr(d, "_tb_metadata_block", None)
    if not block:
        return
    d.add_paragraph("Metadata", style="IndexHeading")
    for line in block.split("\n"):
        d.add_paragraph(line, style="IndexEntry")


# --------------------------------------------------------------------------
# Entry / sub-entry / synthetic / table renderers
# --------------------------------------------------------------------------


def _render_entry(
    d: Document,
    entry: IndexEntry,
    pairs: list[tuple[str, str]] | None = None,
    keep_plural_set: frozenset[str] = frozenset(),
    dropped_plural_acc: list[tuple[str, str]] | None = None,
    removal_set: frozenset[str] = frozenset(),
) -> None:
    """One main paragraph for the entry; one IndexSubentry per sub-entry.

    Phase 7 Wave 3: recapitalize ``pairs`` (CUR-02) applied at the
    string level BEFORE every ``add_run(text)`` (Pitfall 3 of 07-RESEARCH —
    DOCX run-segmentation can never fragment a string-level replace).
    Plural variants (CUR-03) are dropped at ``_filter_variants``.

    Phase 7 Wave 4 (CUR-01 sub-entry inheritance): sub-entries whose
    surface text matches a curator-removed canonical are dropped (RESEARCH
    §Pitfall 4). Default empty ``removal_set`` preserves v1.0 behavior.
    """
    p = d.add_paragraph(style="IndexEntry")
    p.add_run(_recap(entry.canonical, pairs))

    variants = _filter_variants(
        entry.canonical, entry.variants, keep_plural_set, dropped_plural_acc
    )
    if variants:
        also_text = " *(also: " + _VARIANT_SEP.join(variants) + ")*"
        italic = p.add_run(_recap(also_text, pairs))
        italic.italic = True

    locator_str = _render_locators_string(entry.locators)
    if locator_str:
        p.add_run(_recap(f", {locator_str}", pairs))

    for sub in sorted(entry.sub_entries, key=lambda s: s.sort_key):
        if removal_set and is_removed(sub.text, removal_set):
            continue
        sub_p = d.add_paragraph(style="IndexSubentry")
        sub_p.add_run(_recap(sub.text, pairs))
        sub_loc = _render_locators_string(sub.locators)
        if sub_loc:
            sub_p.add_run(_recap(f", {sub_loc}", pairs))


def _render_synthetic(
    d: Document,
    synth: SyntheticEntry,
    entries_by_canonical: dict[str, IndexEntry],
    pairs: list[tuple[str, str]] | None = None,
    removal_set: frozenset[str] = frozenset(),
    transferred_variants: dict[str, tuple[str, ...]] | None = None,
) -> None:
    """Synthetic stem header (IndexEntry) + sibling lines (IndexSubentry).

    CUR-03 sub-rule b (CONTEXT D-13): the literal ``(synthesized)`` marker
    is NOT emitted on the parent line. The IR field
    ``IndexEntry.synthesized`` (and ``SyntheticEntry`` itself) is unchanged;
    this is render-time projection only.

    CUR-01 sub-rule c (RESEARCH §Pitfall 4 + Wave 4 invariant
    test_remove_terms_absent_from_both_outputs): sibling canonicals that
    match a curator-removed term are dropped — they would otherwise leak
    through the top-level is_removed filter (the sibling renders AS A
    SUB-ENTRY of the synthetic stem).

    v1.2.2: ``transferred_variants`` carries the dedup pass's side-channel
    map ``{child_canonical: variants}``. When a same-first-word standalone
    is dropped, its ``*(also: …)*`` parenthetical is re-emitted INLINE on
    this synth's surviving child line (italic run, mirrors the standalone
    entry's ``_render_entry`` shape).
    """
    transferred_variants = transferred_variants or {}
    head = d.add_paragraph(style="IndexEntry")
    head.add_run(_recap(synth.stem, pairs))

    for sib_canonical in sorted(synth.sibling_canonicals):
        if removal_set and is_removed(sib_canonical, removal_set):
            continue
        sib_entry = entries_by_canonical.get(sib_canonical)
        if sib_entry is not None:
            loc_str = _render_locators_string(sib_entry.locators)
        else:
            loc_str = _render_locators_string(list(synth.locators))
        sub_p = d.add_paragraph(style="IndexSubentry")
        sub_p.add_run(_recap(sib_canonical, pairs))
        # v1.2.2: optional transferred variants — italic (also: ...) run
        # mirroring the standalone entry's _render_entry rendering.
        variants = transferred_variants.get(sib_canonical, ())
        if variants:
            filtered = _filter_variants(sib_canonical, list(variants))
            if filtered:
                also_text = " *(also: " + _VARIANT_SEP.join(filtered) + ")*"
                italic = sub_p.add_run(_recap(also_text, pairs))
                italic.italic = True
        if loc_str:
            sub_p.add_run(_recap(f", {loc_str}", pairs))


def _entry_letter(sort_key: str) -> str:
    return sort_key[0].upper() if sort_key else ""


def _render_cross_ref(
    d: Document,
    xref: CrossRefEntry,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    """UAT 08-1: emit ``head. See primary canonical.`` as an IndexEntry
    paragraph (no locators).

    Mirrors the markdown renderer's ``_render_cross_ref_line`` shape so
    both output variants present the same alphabetical anchor to the
    reader. The cross-ref text is recapitalized via ``_recap`` for parity
    with the rest of the entry stream (CUR-02).
    """
    p = d.add_paragraph(style="IndexEntry")
    text = f"{xref.head}. See {xref.primary_canonical}."
    p.add_run(_recap(text, pairs))


def _render_consolidated(
    d: Document,
    consolidated: ConsolidatedEntry,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    """v1.2.1: emit a ConsolidatedEntry as an IndexEntry paragraph.

    For ``source_kind == "xref"``: ``term(s). See target.``
    For ``source_kind == "primary"``: ``term(s), § N.NN (p. N), …``

    Recapitalize is applied at the STRING level BEFORE every
    ``add_run(text)`` (Pitfall 3 of 07-RESEARCH — DOCX run-segmentation
    cannot fragment a string-level replace).
    """
    p = d.add_paragraph(style="IndexEntry")
    if consolidated.source_kind == "xref":
        text = (
            f"{consolidated.display_canonical}. "
            f"See {consolidated.see_target}."
        )
        p.add_run(_recap(text, pairs))
        return
    # primary
    p.add_run(_recap(consolidated.display_canonical, pairs))
    locator_str = _render_locators_string(list(consolidated.locators))  # type: ignore[arg-type]
    if locator_str:
        p.add_run(_recap(f", {locator_str}", pairs))


def _render_subject_index(
    d: Document,
    tree: IndexTree,
    synthetics: list[SyntheticEntry],
    pairs: list[tuple[str, str]] | None = None,
    removal_set: frozenset[str] = frozenset(),
    keep_plural_set: frozenset[str] = frozenset(),
    dropped_plural_acc: list[tuple[str, str]] | None = None,
    editorial_overrides: EditorialOverrides | None = None,
    allow_stale_overrides: bool = False,
) -> None:
    """# Subject Index header + letter dividers + entries + synthetics.

    Phase 7 Wave 3: filters out ``removal_set`` entries (CUR-01); threads
    ``keep_plural_set`` + ``dropped_plural_acc`` to the variant filter
    (CUR-03); recapitalize ``pairs`` flow into every ``add_run`` callsite
    in the entry/synthesized renderers.
    """
    d.add_paragraph(_recap("Subject Index", pairs), style="IndexHeading")

    surviving = [
        e
        for e in tree.entries
        if not is_cruft(e.canonical) and not is_removed(e.canonical, removal_set)
    ]

    # Phase 9 — 7th projection pass (mirrors markdown.py / docx_sections_only.py).
    _xref_removal_set: frozenset[str] = frozenset()
    _synth_suppressed_stems: frozenset[str] = frozenset()
    if editorial_overrides is not None:
        _apply_result = apply_editorial_overrides(
            surviving,
            editorial_overrides,
            allow_stale=allow_stale_overrides,
        )
        surviving = _apply_result.entries
        _xref_removal_set = _apply_result.xref_removal_set
        _synth_suppressed_stems = _apply_result.synth_suppressed_stems
        write_mismatch_report(
            _apply_result.mismatches,
            Path("artifacts/audit/editorial_override_mismatches.md"),
        )

    entries_by_canonical = {e.canonical: e for e in surviving}

    # UAT 08-1 head-noun cross-refs (mirrors markdown.py). Phase 9 R6
    # filters synth_suppressed_stems; R5 filters cross-refs below.
    surviving_synthetics = [
        s for s in synthetics
        if not is_removed(s.stem, removal_set)
        and s.stem not in _synth_suppressed_stems
    ]
    cross_refs = derive_cross_refs(surviving, surviving_synthetics)
    if _xref_removal_set:
        cross_refs = [x for x in cross_refs if x.head not in _xref_removal_set]

    sorted_entries = sorted(surviving, key=lambda e: e.sort_key)
    sorted_synthetics = sorted(surviving_synthetics, key=lambda s: s.stem)

    merged: list[tuple[str, str, object]] = []
    for e in sorted_entries:
        merged.append((e.sort_key, "entry", e))
    for s in sorted_synthetics:
        merged.append((s.stem.lower(), "synth", s))
    for x in cross_refs:
        merged.append((x.sort_key, "xref", x))
    merged.sort(key=lambda t: t[0])

    # v1.2.1: collapse adjacent singular/plural pairs (mirrors markdown.py).
    merged = consolidate_plural_pairs(merged, keep_plural_set)

    # v1.2.2: parent-aliased-standalone dedup (mirrors markdown.py).
    merged, transferred_variants, _dedup_decisions = (
        dedupe_parent_aliased_standalones(merged, entries_by_canonical)
    )

    current_letter = ""
    for sort_key, kind, payload in merged:
        letter = _entry_letter(sort_key)
        if letter != current_letter:
            current_letter = letter
            # Letter dividers are single-letter — len < 2 — so the LLM
            # filters them out of recap pairs (Field min_length=2). Skip
            # _recap defensively to make the no-recap path identical.
            d.add_paragraph(letter, style="IndexHeading")
        if kind == "entry":
            _render_entry(
                d,
                payload,  # type: ignore[arg-type]
                pairs,
                keep_plural_set,
                dropped_plural_acc,
                removal_set,
            )
        elif kind == "synth":
            _render_synthetic(
                d,
                payload,  # type: ignore[arg-type]
                entries_by_canonical,
                pairs,
                removal_set,
                transferred_variants,
            )
        elif kind == "consolidated":
            _render_consolidated(d, payload, pairs)  # type: ignore[arg-type]
        else:  # xref — UAT 08-1
            _render_cross_ref(d, payload, pairs)  # type: ignore[arg-type]


def _render_table_of_cases(
    d: Document,
    tables: TableOfCases | None,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    d.add_paragraph(_recap("Table of Cases", pairs), style="IndexHeading")
    if tables is None:
        return
    for case in sorted(tables.entries, key=lambda c: c.sort_key):
        p = d.add_paragraph(style="IndexEntry")
        # Italic display_name run (TAB-01).
        name_run = p.add_run(_recap(case.display_name, pairs))
        name_run.italic = True
        p.add_run(_recap(f", {case.canonical_citation}", pairs))
        loc_str = _render_locators_string(case.locators)
        if loc_str:
            p.add_run(_recap(f", {loc_str}", pairs))


def _render_table_of_statutes(
    d: Document,
    tables: TableOfStatutes | None,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    d.add_paragraph(_recap("Table of Statutes", pairs), style="IndexHeading")
    if tables is None:
        return
    # B-11: collapse newline-duplicate canonicals at render time (CONTEXT
    # D-02 — render-path only; Phase 3b IR untouched). Helper imported from
    # markdown.py per Option B in Plan 06-03 (single source of truth).
    deduped = _dedup_statute_entries(tables.entries)
    for s in sorted(deduped, key=lambda x: x.sort_key):
        p = d.add_paragraph(style="IndexEntry")
        p.add_run(_recap(s.display_name, pairs))
        loc_str = _render_locators_string(s.locators)
        if loc_str:
            p.add_run(_recap(f", {loc_str}", pairs))


def _render_table_of_rules(
    d: Document,
    tables: TableOfRules | None,
    pairs: list[tuple[str, str]] | None = None,
) -> None:
    d.add_paragraph(_recap("Table of Rules", pairs), style="IndexHeading")
    if tables is None:
        return
    for rule in sorted(tables.entries, key=lambda r: r.sort_key):
        p = d.add_paragraph(style="IndexEntry")
        p.add_run(_recap(rule.parent_rule, pairs))
        parent_loc = _render_locators_string(rule.parent_locators)
        if parent_loc:
            p.add_run(_recap(f", {parent_loc}", pairs))
        for sub in sorted(rule.subsections, key=lambda s: s.subsection_path):
            sub_p = d.add_paragraph(style="IndexSubentry")
            sub_p.add_run(_recap(f"{rule.parent_rule}{sub.subsection_path}", pairs))
            sub_loc = _render_locators_string(sub.locators)
            if sub_loc:
                sub_p.add_run(_recap(f", {sub_loc}", pairs))


# --------------------------------------------------------------------------
# Public API
# --------------------------------------------------------------------------


def render_docx(
    tree: IndexTree,
    synthetics: list[SyntheticEntry],
    tables: dict[str, object],
    metadata: Metadata,
    out_path: Path,
    overrides: CuratorOverrides | None = None,
    editorial_overrides: EditorialOverrides | None = None,
    curator_log: dict[str, list] | None = None,
) -> None:
    """Render OUT-02 — write index.docx to ``out_path`` (always frozen).

    Phase 7 Wave 3: when ``overrides`` is provided, applies the full
    curator pass (CUR-01 removal + CUR-02 recapitalize + CUR-03
    plural-variant filter + (synthesized)-marker drop). Recapitalize is
    applied at the STRING level BEFORE every ``add_run(text)`` call
    (Pitfall 3 — DOCX run-segmentation cannot fragment a string-level
    replace). Audit ledger faithfulness is preserved (Phase 4 IR +
    index_evidence.json untouched).

    Pipeline:
      1. Validate recap pairs (strict-guard re-run for defense in depth).
      2. Build python-docx Document in-memory (Path A — baked paragraphs).
      3. Save to a temp file.
      4. Apply RESEARCH §H-3 ``freeze_docx`` → ``out_path``.
      5. Delete temp.

    Args:
        ...same as v1.0 plus:
        overrides: optional curator overrides; ``None`` reproduces v1.0.
        curator_log: optional accumulator dict (see ``render_markdown``).

    Raises:
        FreezeError: if the freeze post-process fails.
        RecapitalizeGuardError: if any recap pair fails the strict guard.
    """
    out_path = Path(out_path)
    cases_table = tables.get("cases")  # type: ignore[union-attr]
    statutes_table = tables.get("statutes")  # type: ignore[union-attr]
    rules_table = tables.get("rules")  # type: ignore[union-attr]

    pairs: list[tuple[str, str]] | None = None
    removal_set: frozenset[str] = frozenset()
    keep_plural_set: frozenset[str] = frozenset()
    dropped_plural_acc: list[tuple[str, str]] = []

    if overrides is not None:
        removal_set = overrides.removal_set
        keep_plural_set = overrides.keep_plural_set
        if overrides.recapitalize_pairs:
            pairs = [(p[0], p[1]) for p in overrides.recapitalize_pairs]
            assert_letters_only(pairs)

    # Phase 9 — D-07 ALLOW_STALE_OVERRIDES env-flag (mirrors markdown.py).
    allow_stale_overrides = (
        os.environ.get("ALLOW_STALE_OVERRIDES") == "1"
    )

    d = _new_docx_document()
    register_index_styles(d)
    _render_metadata_into_doc(d, metadata)

    _render_subject_index(
        d, tree, synthetics, pairs, removal_set, keep_plural_set, dropped_plural_acc,
        editorial_overrides=editorial_overrides,
        allow_stale_overrides=allow_stale_overrides,
    )
    _render_table_of_cases(d, cases_table, pairs)  # type: ignore[arg-type]
    _render_table_of_statutes(d, statutes_table, pairs)  # type: ignore[arg-type]
    _render_table_of_rules(d, rules_table, pairs)  # type: ignore[arg-type]
    _emit_metadata_block_paragraph(d)

    # Phase 9 — R9 whitespace at run-text level. Walk every paragraph's
    # runs and apply the post-emit string replace. Mirrors the markdown
    # post-emit hook but operates per-run because docx text lives in
    # python-docx run objects, not a single string.
    if editorial_overrides is not None and editorial_overrides.R9_whitespace:
        r9_rules = editorial_overrides.R9_whitespace
        for paragraph in d.paragraphs:
            for run in paragraph.runs:
                run.text = apply_r9_whitespace(run.text, r9_rules)

    if curator_log is not None and overrides is not None:
        curator_log.setdefault("dropped_plural_variants", []).extend(
            dropped_plural_acc
        )
        # Note: dangling-xref auto-strip is markdown-only in this slice
        # (the v1.0 DOCX renderer does not emit `See <target>` clauses
        # outside the IR-driven entry/sub-entry shape, and `(also: …)`
        # parentheticals already use the curator-aware _filter_variants
        # path which short-circuits removed targets via the underlying
        # IR-side `is_removed` check). The markdown audit log is the
        # single source of truth for dangling_xrefs_stripped.

    # Save to temp, then freeze to out_path.
    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, dir=out_path.parent
    ) as tmp_fp:
        tmp_path = Path(tmp_fp.name)
    try:
        d.save(str(tmp_path))
        freeze_docx(tmp_path, out_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()
